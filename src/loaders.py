from pathlib import Path

import extract_msg
import pandas as pd
from docx import Document as WordDocument
from langchain_community.document_loaders import UnstructuredPDFLoader
from langchain_core.documents import Document
from openpyxl import load_workbook


def documents_list(folder_path, extensions=[".docx", ".pdf", ".xlsx"]):
    """
    Scans a folder recursively and returns a dictionary of file paths grouped by extension.
    Args:
        folder_path (str or Path): The path to the folder to scan.
        extensions (list of str, optional): List of file extensions to look for. Defaults to [".docx", ".pdf", ".xlsx"].
    Returns:
        dict: A dictionary where keys are file extensions and values are lists of file paths matching that extension.
    """
    documents_dict = {ext: [] for ext in extensions}
    for key in documents_dict.keys():
        documents_dict[key].extend(
            str(file) for file in Path(folder_path).rglob(f"*{key}")
        )
    return documents_dict


def load_table(table):
    """
    Converts a Word table into a markdown-formatted string.
    Args:
        table (docx.table.Table): A table object from a Word document.
    Returns:
        str: The table converted to a markdown-style string.
    """
    markdown = []
    for row in table.rows:
        cell = [cell.text.strip() for cell in row.cells]
        markdown.append("|" + "|".join(cell) + "|")
    return "\n".join(markdown)


def load_word(file_path):
    """
    Loads a Word document, extracting paragraphs and tables into Document objects.
    Args:
        file_path (str): Path to the Word (.docx) file.
    Returns:
        list[Document]: A list of Document objects containing paragraphs and tables with metadata including the source file path.
    """
    document = WordDocument(file_path)

    document_object = []
    para_text = ""
    for para in document.paragraphs:

        text = para.text.strip()
        if text:
            para_text += text + "\n"
    document_object.append(
        Document(page_content=para_text, metadata={"source": file_path})
    )
    for table in document.tables:
        document_object.append(
            Document(page_content=load_table(table), metadata={"source": file_path})
        )
    return document_object


def load_excel(file_path):
    """
    Loads an Excel file, extracting each sheet into Document objects.
    Args:
        file_path (str): Path to the Excel (.xlsx) file.
    Returns:
        list[Document]: A list of Document objects for each sheet, with rows formatted as key-value pairs and metadata including the source file and sheet name.
    """
    xls = pd.ExcelFile(file_path)
    sheet_names = xls.sheet_names
    document_objects = []
    for sheet in sheet_names:
        df = pd.read_excel(file_path, sheet_name=sheet)
        header = list(df.columns)
        document_data = []

        for _, row in df.iterrows():
            row_data = []
            for col in header:
                row_data.append(f"{col}: {row[col]}")
            document_data.append(" | ".join(row_data))
        document_objects.append(
            Document(
                page_content="\n".join(document_data),
                metadata={"source": file_path, "sheet": sheet},
            )
        )
    return document_objects


def load_email(file_path):
    """
    Loads an email file (.msg) and extracts the body and metadata.
    Args:
        file_path (str): Path to the email (.msg) file.
    Returns:
        list[Document]: A list containing a single Document object with the email body and metadata including subject, sender, date, and source.
    """
    msg = extract_msg(file_path)
    if len(msg.body) > 5:
        doc = Document(
            page_content=msg.body,
            metadata={
                "subject": msg.subject,
                "sender": msg.sender,
                "date": msg.date,
                "source": file_path,
            },
        )
        return [doc]


def load_text(file_path):
    """
    Loads a plain text file into a Document object.
    Args:
        file_path (str): Path to the text (.txt) file.
    Returns:
        list[Document]: A list containing a single Document object with the file content and source metadata.
    """
    with open(file_path, "r") as file:
        text = file.read()
    return [Document(page_content=text, metadata={"source": file_path})]


def load_pdf(file_path):
    """
    Loads a PDF file using UnstructuredPDFLoader and returns its content as Document objects.
    Args:
        file_path (str): Path to the PDF file.
    Returns:
        list[Document]: A list of Document objects containing the PDF content.
    """
    loader = UnstructuredPDFLoader(file_path)
    return loader.load()
